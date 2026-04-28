"""
Table support test for docx_writer.
"""

import unittest
from pathlib import Path
from docx import Document

from src.docx_writer import DocxWriter


class TestTablePlaceholderSupport(unittest.TestCase):
    """Test placeholders inside tables work correctly."""

    @classmethod
    def setUpClass(cls):
        cls.test_dir = Path("tests/test_data")
        cls.test_dir.mkdir(parents=True, exist_ok=True)

        cls.test_template = cls.test_dir / "table_test.docx"
        doc = Document()
        
        # Paragraph outside table
        doc.add_paragraph("Before table: {{OUTER}}")
        
        # Add table with 1 row, 2 cells
        table = doc.add_table(rows=1, cols=2)
        
        # Cell 0: placeholder
        cell0 = table.cell(0, 0)
        cell0.text = "Cell content: {{CELL_PLACEHOLDER}}"
        
        # Cell 1: mixed
        cell1_para = table.cell(0, 1).add_paragraph()
        cell1_para.add_run("Start {{INLINE}} end")
        
        # Paragraph after table
        doc.add_paragraph("After table: {{ANOTHER}}")
        
        doc.save(cls.test_template)

    def test_table_placeholders_detected(self):
        """Test that placeholders inside table cells are found."""
        writer = DocxWriter(self.test_template)
        sections = writer.get_section_names()
        self.assertIn("CELL_PLACEHOLDER", sections)
        self.assertIn("INLINE", sections)
        self.assertIn("OUTER", sections)
        self.assertIn("ANOTHER", sections)

    def test_table_placeholder_replacement(self):
        """Test replacing placeholders inside table cells."""
        writer = DocxWriter(self.test_template)
        replacements = {
            "CELL_PLACEHOLDER": "REPLACED_CELL",
            "INLINE": "INLINE_VAL",
            "OUTER": "OUTER_VAL",
            "ANOTHER": "ANOTHER_VAL"
        }
        count = writer.replace_placeholders(replacements)
        self.assertEqual(count, 4)

        # Save and re-read to verify persistence
        output_path = self.test_dir / "table_output.docx"
        writer.save(output_path)

        doc2 = Document(output_path)
        
        # Check body paragraphs
        body_paras = [p.text for p in doc2.paragraphs]
        combined_body = "\n".join(body_paras)
        self.assertIn("OUTER_VAL", combined_body)
        self.assertIn("ANOTHER_VAL", combined_body)
        
        # Check table cells
        found_cell = False
        found_inline = False
        for table in doc2.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if "REPLACED_CELL" in cell_text:
                        found_cell = True
                    if "INLINE_VAL" in cell_text:
                        found_inline = True
        
        self.assertTrue(found_cell, "Cell placeholder replacement not found in any table cell")
        self.assertTrue(found_inline, "Inline placeholder replacement not found in any table cell")

    def test_nested_table_placeholders(self):
        """Test placeholders in nested tables."""
        doc = Document()
        table1 = doc.add_table(rows=1, cols=1)
        cell = table1.cell(0, 0)
        table2 = cell.add_table(rows=1, cols=1)
        nested_cell = table2.cell(0, 0)
        nested_cell.text = "Nested: {{NESTED_PLACEHOLDER}}"
        
        nested_path = self.test_dir / "nested_table.docx"
        doc.save(nested_path)
        
        writer = DocxWriter(nested_path)
        self.assertIn("NESTED_PLACEHOLDER", writer.get_section_names())
        
        writer.replace_placeholders({"NESTED_PLACEHOLDER": "NESTED_OK"})
        
        # Verify by reading back using known structure
        output_path = self.test_dir / "nested_output.docx"
        writer.save(output_path)
        doc3 = Document(output_path)
        
        # Navigate to nested cell: top-level table -> outer cell -> nested table -> inner cell
        outer_table = doc3.tables[0]
        outer_cell = outer_table.cell(0, 0)
        self.assertTrue(len(outer_cell.tables) > 0, "Nested table missing in outer cell")
        inner_table = outer_cell.tables[0]
        inner_cell = inner_table.cell(0, 0)
        self.assertIn("NESTED_OK", inner_cell.text)


if __name__ == "__main__":
    unittest.main()
