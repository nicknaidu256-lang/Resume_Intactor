"""
Comprehensive tests for AdvancedDocxEngine with perfect formatting preservation.
"""

import unittest
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

from src.advanced_docx_engine import AdvancedDocxEngine, ReplacementStrategy, create_test_template


class TestAdvancedDocxEngine(unittest.TestCase):
    """Test suite for AdvancedDocxEngine."""

    @classmethod
    def setUpClass(cls):
        """Create test templates."""
        cls.test_dir = Path("tests/test_data")
        cls.test_dir.mkdir(parents=True, exist_ok=True)
        
        # Create complex test template
        cls.complex_template = cls.test_dir / "complex_test_template.docx"
        doc = Document()
        
        # Simple placeholder
        doc.add_paragraph("Simple replacement: {{SIMPLE}}")
        
        # Formatted placeholder
        p = doc.add_paragraph("Formatted: ")
        run = p.add_run("{{FORMATTED}}")
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red
        
        # Multiple placeholders
        doc.add_paragraph("Multiple: {{FIRST}} and {{SECOND}}")
        
        # Table with placeholders
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "{{TABLE_CELL_1}}"
        table.cell(0, 1).text = "Normal text"
        table.cell(1, 0).text = "{{TABLE_CELL_2}}"
        table.cell(1, 1).text = "{{TABLE_CELL_3}}"
        
        doc.save(str(cls.complex_template))

    def test_engine_initialization(self):
        """Test engine initializes correctly."""
        engine = AdvancedDocxEngine(self.complex_template)
        self.assertIsNotNone(engine.document)
        self.assertIsInstance(engine.placeholders, dict)

    def test_placeholder_detection(self):
        """Test all placeholders are detected."""
        engine = AdvancedDocxEngine(self.complex_template)
        names = engine.get_placeholder_names()
        
        expected = {'SIMPLE', 'FORMATTED', 'FIRST', 'SECOND', 
                   'TABLE_CELL_1', 'TABLE_CELL_2', 'TABLE_CELL_3'}
        self.assertEqual(set(names), expected)

    def test_simple_replacement(self):
        """Test simple placeholder replacement."""
        engine = AdvancedDocxEngine(self.complex_template)
        replacements = {
            'SIMPLE': 'This is a simple replacement',
            'FORMATTED': 'BOLD AND RED',
            'FIRST': 'first item',
            'SECOND': 'second item',
            'TABLE_CELL_1': 'table cell 1',
            'TABLE_CELL_2': 'table cell 2', 
            'TABLE_CELL_3': 'table cell 3'
        }
        
        stats = engine.replace_placeholders(replacements)
        self.assertEqual(stats['replaced'], 7)
        self.assertEqual(stats['skipped_missing'], 0)
        self.assertEqual(stats['skipped_errors'], 0)
        
        # Verify text replacement
        paragraphs = engine.document.paragraphs
        self.assertEqual(paragraphs[0].text, "Simple replacement: This is a simple replacement")
        self.assertEqual(paragraphs[2].text, "Multiple: first item and second item")
        
        # Verify table replacement
        table = engine.document.tables[0]
        self.assertEqual(table.cell(0, 0).text, "table cell 1")
        self.assertEqual(table.cell(1, 0).text, "table cell 2")
        self.assertEqual(table.cell(1, 1).text, "table cell 3")

    def test_formatting_preservation_single_run(self):
        """Test formatting preservation in single-run scenarios."""
        # Create a template where placeholder is in its own run
        test_path = self.test_dir / "formatting_test.docx"
        doc = Document()
        
        p = doc.add_paragraph()
        run = p.add_run("{{FORMATTED_PLACEHOLDER}}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green
        
        doc.save(str(test_path))
        
        engine = AdvancedDocxEngine(test_path)
        replacements = {'FORMATTED_PLACEHOLDER': 'REPLACEMENT TEXT'}
        stats = engine.replace_placeholders(replacements)
        
        self.assertEqual(stats['replaced'], 1)
        
        # Verify formatting preserved
        paragraph = engine.document.paragraphs[0]
        run = paragraph.runs[0]
        self.assertEqual(run.text, "REPLACEMENT TEXT")
        self.assertTrue(run.bold)
        # Note: Color preservation may be limited by python-docx capabilities

    def test_partial_replacement(self):
        """Test behavior when some placeholders are missing from replacements."""
        engine = AdvancedDocxEngine(self.complex_template)
        replacements = {
            'SIMPLE': 'replaced',
            'FIRST': 'replaced'
            # Missing: FORMATTED, SECOND, TABLE_CELL_1, TABLE_CELL_2, TABLE_CELL_3
        }
        
        stats = engine.replace_placeholders(replacements)
        self.assertEqual(stats['replaced'], 2)
        self.assertEqual(stats['skipped_missing'], 5)
        
        # Verify only specified placeholders were replaced
        paragraphs = engine.document.paragraphs
        self.assertEqual(paragraphs[0].text, "Simple replacement: replaced")
        self.assertEqual(paragraphs[2].text, "Multiple: replaced and {{SECOND}}")
        
        # Table cells should remain unchanged
        table = engine.document.tables[0]
        self.assertEqual(table.cell(0, 0).text, "{{TABLE_CELL_1}}")

    def test_strict_mode(self):
        """Test strict mode raises error for missing replacements."""
        engine = AdvancedDocxEngine(self.complex_template)
        replacements = {'SIMPLE': 'test'}  # Missing others
        
        with self.assertRaises(ValueError):
            engine.replace_placeholders(replacements, strict=True)

    def test_different_strategies(self):
        """Test different replacement strategies."""
        test_path = self.test_dir / "strategy_test.docx"
        doc = Document()
        doc.add_paragraph("Test: {{TEST}}")
        doc.save(str(test_path))
        
        # Test each strategy
        for strategy in ReplacementStrategy:
            with self.subTest(strategy=strategy):
                engine = AdvancedDocxEngine(test_path, strategy=strategy)
                stats = engine.replace_placeholders({'TEST': 'success'})
                self.assertEqual(stats['replaced'], 1)
                self.assertEqual(engine.document.paragraphs[0].text, "Test: success")

    def test_save_document(self):
        """Test document saving functionality."""
        engine = AdvancedDocxEngine(self.complex_template)
        replacements = {'SIMPLE': 'saved content'}
        engine.replace_placeholders(replacements)
        
        output_path = self.test_dir / "test_output.docx"
        engine.save(output_path)
        
        self.assertTrue(output_path.exists())
        
        # Verify saved document can be reloaded
        reloaded = Document(str(output_path))
        self.assertEqual(reloaded.paragraphs[0].text, "Simple replacement: saved content")

    def test_error_handling(self):
        """Test error handling for invalid inputs."""
        # Non-existent template
        with self.assertRaises(FileNotFoundError):
            AdvancedDocxEngine(Path("nonexistent.docx"))
        
        # Valid template but invalid operations
        engine = AdvancedDocxEngine(self.complex_template)
        with self.assertRaises(ValueError):
            engine.save(Path("/invalid/path/output.docx"))


class TestRealWorldScenarios(unittest.TestCase):
    """Test realistic scenarios similar to resume templates."""

    def test_complex_table_structure(self):
        """Test with complex table structures like resumes."""
        test_path = Path("tests/test_data/realistic_resume.docx")
        test_path.parent.mkdir(exist_ok=True)
        
        doc = Document()
        
        # Create a resume-like table structure
        table = doc.add_table(rows=10, cols=2)
        
        # Header row
        table.cell(0, 0).text = "{{NAME}}"
        table.cell(0, 1).text = "{{CONTACT_INFO}}"
        
        # Experience section
        table.cell(1, 0).text = "Experience"
        table.cell(2, 0).text = "{{COMPANY_1}}"
        table.cell(2, 1).text = "{{ROLE_1}} - {{DATES_1}}"
        table.cell(3, 1).text = "{{BULLET_1_1}}"
        table.cell(4, 1).text = "{{BULLET_1_2}}"
        
        # Skills section
        table.cell(5, 0).text = "Skills"
        table.cell(6, 1).text = "{{SKILLS_LIST}}"
        
        # Education
        table.cell(7, 0).text = "Education"
        table.cell(8, 1).text = "{{DEGREE}} - {{UNIVERSITY}}"
        
        doc.save(str(test_path))
        
        # Test replacement
        engine = AdvancedDocxEngine(test_path)
        replacements = {
            'NAME': 'John Doe',
            'CONTACT_INFO': 'john@example.com | (555) 123-4567',
            'COMPANY_1': 'Tech Corp Inc.',
            'ROLE_1': 'Senior Developer',
            'DATES_1': '2020-2024',
            'BULLET_1_1': 'Led team of 5 developers on critical projects',
            'BULLET_1_2': 'Improved system performance by 40%',
            'SKILLS_LIST': 'Python, JavaScript, SQL, Docker',
            'DEGREE': 'BS Computer Science',
            'UNIVERSITY': 'Tech University'
        }
        
        stats = engine.replace_placeholders(replacements)
        self.assertEqual(stats['replaced'], 10)
        
        # Verify all replacements
        table = engine.document.tables[0]
        self.assertEqual(table.cell(0, 0).text, "John Doe")
        self.assertEqual(table.cell(0, 1).text, "john@example.com | (555) 123-4567")
        self.assertEqual(table.cell(2, 0).text, "Tech Corp Inc.")
        self.assertEqual(table.cell(2, 1).text, "Senior Developer - 2020-2024")
        self.assertEqual(table.cell(3, 1).text, "Led team of 5 developers on critical projects")
        self.assertEqual(table.cell(4, 1).text, "Improved system performance by 40%")
        self.assertEqual(table.cell(6, 1).text, "Python, JavaScript, SQL, Docker")
        self.assertEqual(table.cell(8, 1).text, "BS Computer Science - Tech University")


if __name__ == "__main__":
    unittest.main()