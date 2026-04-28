"""
Integration tests for SafeDocxEngine with real-world scenarios.
"""

import unittest
from pathlib import Path
from docx import Document

from src.safe_docx_engine import SafeDocxEngine


class TestSafeDocxIntegration(unittest.TestCase):
    """Integration tests with realistic resume templates."""

    def test_full_resume_replacement(self):
        """Test complete resume template replacement."""
        engine = SafeDocxEngine("templates/Master_Resume.docx")
        
        # Comprehensive replacements matching the actual template structure
        replacements = {
            'SUMMARY': 'Senior Software Engineer with 8+ years of experience in full-stack development. '
                      'Specialized in building scalable web applications using modern technologies.',
            'EXP1_BULLET1': 'Led development of microservices architecture serving 2M+ users',
            'EXP1_BULLET2': 'Managed team of 6 developers using Agile methodologies',
            'EXP1_BULLET3': 'Implemented CI/CD pipelines reducing deployment time by 70%',
            'EXP1_BULLET4': 'Optimized database queries improving response time by 45%',
            'EXP1_BULLET5': 'Designed and developed RESTful APIs with 99.9% uptime',
            'EXP1_BULLET6': 'Mentored junior developers and conducted code reviews',
            'EXP1_BULLET7': 'Integrated third-party services and payment gateways',
            'EXP1_BULLET8': 'Implemented comprehensive testing strategy with 95% coverage',
            'EXP1_BULLET9': 'Collaborated with product teams on feature planning and delivery',
            'SKILLS_1': 'Programming: Python, JavaScript, Java, TypeScript, SQL',
            'SKILLS_2': 'Frameworks: Django, React, Spring Boot, Node.js',
            'SKILLS_3': 'Tools: Docker, Kubernetes, AWS, Git, Jenkins, PostgreSQL'
        }
        
        stats = engine.replace_placeholders(replacements)
        
        # Should replace all 13 placeholders
        self.assertEqual(stats['replaced'], 13)
        self.assertEqual(stats['skipped_missing'], 0)
        self.assertEqual(stats['skipped_errors'], 0)
        
        # Save and verify
        output_path = Path("tests/test_data/full_replacement_test.docx")
        engine.save(output_path)
        
        # Verify file was created
        self.assertTrue(output_path.exists())
        
        # Reload and verify content
        reloaded = Document(str(output_path))
        
        # The document should have content (not just placeholders)
        # Since it's a table-based resume, we check that it's not empty
        self.assertTrue(any(cell.text.strip() for row in reloaded.tables[0].rows 
                          for cell in row.cells))

    def test_partial_replacement_preservation(self):
        """Test that unreplaced placeholders are preserved."""
        engine = SafeDocxEngine("templates/Master_Resume.docx")
        
        # Only replace some placeholders
        replacements = {
            'SUMMARY': 'Test summary replacement',
            'EXP1_BULLET1': 'Test bullet point 1',
            'SKILLS_1': 'Test skills list'
        }
        
        stats = engine.replace_placeholders(replacements)
        
        # Should replace 3, skip 10
        self.assertEqual(stats['replaced'], 3)
        self.assertEqual(stats['skipped_missing'], 10)
        
        # Save and verify
        output_path = Path("tests/test_data/partial_replacement_test.docx")
        engine.save(output_path)
        
        # Reload and check that unreplaced placeholders are still there
        reloaded = Document(str(output_path))
        
        # Should contain both replaced content and original placeholders
        content = '\n'.join(cell.text for row in reloaded.tables[0].rows for cell in row.cells)
        self.assertIn('Test summary replacement', content)
        self.assertIn('Test bullet point 1', content)
        self.assertIn('Test skills list', content)
        self.assertIn('{{EXP1_BULLET2}}', content)  # Should remain
        self.assertIn('{{SKILLS_2}}', content)      # Should remain

    def test_empty_replacements(self):
        """Test behavior when no replacements are provided."""
        engine = SafeDocxEngine("templates/Master_Resume.docx")
        
        stats = engine.replace_placeholders({})
        
        # Should skip all placeholders
        self.assertEqual(stats['replaced'], 0)
        self.assertEqual(stats['skipped_missing'], 13)
        
        # Document should remain unchanged
        output_path = Path("tests/test_data/empty_replacement_test.docx")
        engine.save(output_path)
        
        # All placeholders should still be present
        reloaded = Document(str(output_path))
        content = '\n'.join(cell.text for row in reloaded.tables[0].rows for cell in row.cells)
        self.assertIn('{{SUMMARY}}', content)
        self.assertIn('{{EXP1_BULLET1}}', content)
        self.assertIn('{{SKILLS_1}}', content)

    def test_large_content_replacement(self):
        """Test replacement with large content blocks."""
        engine = SafeDocxEngine("templates/Master_Resume.docx")
        
        # Large content that might stress the system
        large_summary = ('Experienced professional with extensive background in software development. '
                       'Specialized in building scalable systems using modern technologies. '
                       'Strong problem-solving skills and ability to work in fast-paced environments. '
                       'Excellent communication skills and team collaboration experience. '
                       'Committed to writing clean, maintainable code and following best practices.')
        
        large_bullet = ('Led development of critical system components serving millions of users. '
                       'Implemented advanced caching strategies and database optimization techniques. '
                       'Collaborated with cross-functional teams including product, design, and QA. '
                       'Mentored junior developers and established coding standards. '
                       'Participated in architecture design and technology selection decisions.')
        
        replacements = {
            'SUMMARY': large_summary,
            'EXP1_BULLET1': large_bullet,
            'SKILLS_1': 'Python, JavaScript, Java, SQL, Docker, Kubernetes, AWS, React, Django, Spring Boot'
        }
        
        stats = engine.replace_placeholders(replacements)
        self.assertEqual(stats['replaced'], 3)
        
        output_path = Path("tests/test_data/large_content_test.docx")
        engine.save(output_path)
        
        # Verify content was inserted
        reloaded = Document(str(output_path))
        content = '\n'.join(cell.text for row in reloaded.tables[0].rows for cell in row.cells)
        self.assertIn('Experienced professional', content)
        self.assertIn('Led development', content)
        self.assertIn('Python, JavaScript', content)

    def test_output_directory_creation(self):
        """Test that output directories are created automatically."""
        engine = SafeDocxEngine("templates/Master_Resume.docx")
        
        # Use a path that doesn't exist yet
        output_path = Path("tests/test_outputs/new_directory/test_resume.docx")
        
        # Directory shouldn't exist initially
        if output_path.parent.exists():
            import shutil
            shutil.rmtree(output_path.parent)
        
        self.assertFalse(output_path.parent.exists())
        
        # Save should create the directory
        engine.save(output_path)
        self.assertTrue(output_path.parent.exists())
        self.assertTrue(output_path.exists())


if __name__ == "__main__":
    unittest.main()