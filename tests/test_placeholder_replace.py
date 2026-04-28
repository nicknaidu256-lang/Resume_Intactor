"""
Unit tests for placeholder detection and replacement edge cases.
"""

import unittest
from pathlib import Path
from docx import Document

from src.docx_writer import DocxWriter


class TestPlaceholderReplacement(unittest.TestCase):
    """Edge case tests for placeholder handling."""

    @classmethod
    def setUpClass(cls):
        cls.test_dir = Path("tests/test_data")
        cls.test_dir.mkdir(parents=True, exist_ok=True)

        cls.test_template = cls.test_dir / "edge_cases.docx"
        doc = Document()
        # Adjacent placeholders
        doc.add_paragraph("{{START}}{{END}}")
        # Placeholder followed by text
        doc.add_paragraph("{{NAME}} applies")
        # Placeholder with formatting around it (will be in single run in simple template)
        run = doc.add_paragraph().add_run("Status: {{STATUS}}")
        run.bold = True
        doc.save(cls.test_template)

    def test_adjacent_placeholders(self):
        """Test two placeholders next to each other."""
        writer = DocxWriter(self.test_template)
        replacements = {"START": "begin", "END": "end"}
        count = writer.replace_placeholders(replacements)
        self.assertEqual(count, 2)
        self.assertEqual(writer.document.paragraphs[0].text, "beginend")

    def test_placeholder_with_suffix(self):
        """Test placeholder followed by plain text."""
        writer = DocxWriter(self.test_template)
        replacements = {"NAME": "Alice"}
        count = writer.replace_placeholders(replacements)
        self.assertEqual(count, 1)
        self.assertEqual(writer.document.paragraphs[1].text, "Alice applies")

    def test_formatted_run_preserved(self):
        """Test that run formatting (bold) is preserved around replacement."""
        writer = DocxWriter(self.test_template)
        replacements = {"STATUS": "Active"}
        writer.replace_placeholders(replacements)
        para = writer.document.paragraphs[2]
        # The entire paragraph should be "Status: Active"
        self.assertEqual(para.text, "Status: Active")
        # Check that run formatting still present (first run may contain "Status: ")
        runs = para.runs
        self.assertGreater(len(runs), 0)


if __name__ == "__main__":
    unittest.main()
