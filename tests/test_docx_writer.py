"""
Unit tests for docx_writer module.
"""

import unittest
from pathlib import Path
from docx import Document

from src.docx_writer import DocxWriter, PLACEHOLDER_PATTERN


class TestDocxWriter(unittest.TestCase):
    """Test suite for DocxWriter."""

    @classmethod
    def setUpClass(cls):
        """Create a test template document."""
        cls.test_dir = Path("tests/test_data")
        cls.test_dir.mkdir(parents=True, exist_ok=True)

        # Create simple test template
        cls.test_template = cls.test_dir / "test_template.docx"
        doc = Document()
        doc.add_paragraph("Hello {{NAME}}, welcome to {{COMPANY}}.")
        doc.add_paragraph("Your role: {{JOB_TITLE}}")
        doc.add_paragraph("No placeholders here.")
        doc.save(cls.test_template)

    def test_placeholder_pattern_matches(self):
        """Test regex pattern identifies placeholders."""
        text = "Hello {{NAME}}, your position is {{JOB_TITLE}}."
        matches = list(PLACEHOLDER_PATTERN.finditer(text))
        self.assertEqual(len(matches), 2)
        self.assertEqual(matches[0].group(1), "NAME")
        self.assertEqual(matches[1].group(1), "JOB_TITLE")

    def test_load_template(self):
        """Test template loads without errors."""
        writer = DocxWriter(self.test_template)
        self.assertIsNotNone(writer.document)
        self.assertEqual(len(writer.document.paragraphs), 3)

    def test_placeholder_scanning(self):
        """Test placeholders are detected correctly."""
        writer = DocxWriter(self.test_template)
        placeholders = writer.get_section_names()
        self.assertIn("NAME", placeholders)
        self.assertIn("COMPANY", placeholders)
        self.assertIn("JOB_TITLE", placeholders)
        self.assertEqual(len(placeholders), 3)

    def test_original_content_retrieval(self):
        """Test original placeholder content is stored."""
        writer = DocxWriter(self.test_template)
        self.assertEqual(writer.get_original_content("NAME"), "{{NAME}}")

    def test_replace_single_placeholder(self):
        """Test replacing a single placeholder."""
        writer = DocxWriter(self.test_template)
        replacements = {
            "NAME": "John Doe",
            "COMPANY": "Acme Corp",
            "JOB_TITLE": "Software Engineer"
        }
        count = writer.replace_placeholders(replacements)
        self.assertEqual(count, 3)  # 3 placeholders replaced

        # Verify paragraph 0 updated
        self.assertEqual(writer.document.paragraphs[0].text,
                         "Hello John Doe, welcome to Acme Corp.")
        self.assertEqual(writer.document.paragraphs[1].text,
                         "Your role: Software Engineer")

    def test_replace_partial_missing(self):
        """Test behavior when some placeholders missing from replacements dict."""
        writer = DocxWriter(self.test_template)
        # Only provide NAME, leave COMPANY blank
        replacements = {"NAME": "John Doe"}
        count = writer.replace_placeholders(replacements)
        self.assertEqual(count, 1)  # Only NAME replaced

        # COMPANY should remain as {{COMPANY}}
        self.assertIn("{{COMPANY}}", writer.document.paragraphs[0].text)
        self.assertIn("John Doe", writer.document.paragraphs[0].text)

    def test_preserve_paragraph_without_placeholder(self):
        """Test that non-placeholder paragraphs remain untouched."""
        writer = DocxWriter(self.test_template)
        replacements = {"NAME": "John"}
        writer.replace_placeholders(replacements)
        self.assertEqual(writer.document.paragraphs[2].text, "No placeholders here.")


if __name__ == "__main__":
    unittest.main()
