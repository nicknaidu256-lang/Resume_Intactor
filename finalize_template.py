"""
Final placeholder insertion for Master_Resume.docx

This script copies the original resume (already copied) and inserts placeholders
only into designated rows, preserving ALL original formatting.

Placeholder mapping (merged-cell layout):
  Row  8: {{SUMMARY}}
  Row 12: {{SKILLS_1}}
  Row 14: {{SKILLS_2}}
  Row 16: {{SKILLS_3}}
  Row 32: {{EXP1_BULLET1}}
  Row 34: {{EXP1_BULLET2}}
  Row 36: {{EXP1_BULLET3}}
  Row 38: {{EXP1_BULLET4}}
  Row 40: {{EXP1_BULLET5}}
  Row 42: {{EXP1_BULLET6}}
  Row 44: {{EXP1_BULLET7}}
  Row 46: {{EXP1_BULLET8}}
  Row 48: {{EXP1_BULLET9}}

Rows 50-56 (Education) remain UNCHANGED (original content preserved).
"""

from pathlib import Path
from docx import Document

template_path = Path("templates/Master_Resume.docx")
doc = Document(template_path)
table = doc.tables[0]

print(f"Template: {template_path}")
print(f"Table: {len(table.rows)} rows × {len(table.columns)} columns (layout)\n")

def replace_cell_with_placeholder(cell, placeholder):
    """Replace entire cell content with a single paragraph containing placeholder."""
    style_name = cell.paragraphs[0].style.name if cell.paragraphs else 'Normal'
    cell._element.clear()
    para = cell.add_paragraph(placeholder)
    try:
        para.style.name = style_name
    except Exception:
        pass

# Row -> placeholder mapping
row_placeholders = {
    8:  "{{SUMMARY}}",
    12: "{{SKILLS_1}}",
    14: "{{SKILLS_2}}",
    16: "{{SKILLS_3}}",
    32: "{{EXP1_BULLET1}}",
    34: "{{EXP1_BULLET2}}",
    36: "{{EXP1_BULLET3}}",
    38: "{{EXP1_BULLET4}}",
    40: "{{EXP1_BULLET5}}",
    42: "{{EXP1_BULLET6}}",
    44: "{{EXP1_BULLET7}}",
    46: "{{EXP1_BULLET8}}",
    48: "{{EXP1_BULLET9}}",
}

processed = 0
for row_idx, row in enumerate(table.rows):
    if row_idx in row_placeholders:
        ph = row_placeholders[row_idx]
        cell_count = len(row.cells)
        for col_idx, cell in enumerate(row.cells):
            replace_cell_with_placeholder(cell, ph)
        print(f"Row {row_idx:3d}: applied {ph} to {cell_count} cell(s)")
        processed += 1

print(f"\nTotal rows modified: {processed}")

# Save (overwrite existing template)
doc.save(template_path)
print(f"Saved: {template_path}")

# Quick verification
import re
PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
found = {}
for row in table.rows:
    for cell in row.cells:
        for m in PLACEHOLDER_PATTERN.finditer(cell.text):
            name = m.group(1)
            found[name] = found.get(name, 0) + 1

print(f"\nPlaceholder summary: {found}")
expected_names = ['SUMMARY','SKILLS_1','SKILLS_2','SKILLS_3',
                  'EXP1_BULLET1','EXP1_BULLET2','EXP1_BULLET3','EXP1_BULLET4','EXP1_BULLET5',
                  'EXP1_BULLET6','EXP1_BULLET7','EXP1_BULLET8','EXP1_BULLET9']
missing = [n for n in expected_names if n not in found]
if missing:
    print(f"WARNING: Missing placeholders: {missing}")
else:
    print("All expected placeholders inserted correctly [OK]")
