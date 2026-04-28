"""
Insert placeholders into template by replacing content of all cells in target rows.
Uses row iteration to avoid indexing issues with merged cells.
"""

from pathlib import Path
from docx import Document
import re

src = Path("Archive/Original_Resume_Master.docx")
dst = Path("templates/Master_Resume.docx")

print(f"Loading: {src}")
doc = Document(src)
table = doc.tables[0]

print(f"Table dimensions: {len(table.rows)} rows × 4 columns (layout)")

def replace_cell_with_placeholder(cell, placeholder):
    """Replace all content in a cell with a single paragraph having the placeholder."""
    style_name = cell.paragraphs[0].style.name if cell.paragraphs else 'Normal'
    cell._element.clear()
    para = cell.add_paragraph(placeholder)
    try:
        para.style.name = style_name
    except Exception:
        pass
    return style_name

# Mapping: row_index -> placeholder text
row_placeholders = {
    8:  "{{SUMMARY}}",
    12: "{{SKILLS_SECTION}}",
    14: "{{SKILLS_SECTION}}",
    16: "{{SKILLS_SECTION}}",
    32: "{{EXP1_BULLET1}}",
    34: "{{EXP1_BULLET2}}",
    36: "{{EXP1_BULLET3}}",
    38: "{{EXP1_BULLET4}}",
    40: "{{EXP1_BULLET5}}",
    42: "{{EXP1_BULLET6}}",
    44: "{{EXP1_BULLET7}}",
    46: "{{EXP1_BULLET8}}",
    48: "{{EXP1_BULLET9}}",
    52: "{{EDUCATION_SECTION}}",
}

# Rows to clear (education details that become part of EDUCATION_SECTION)
clear_rows = {53, 55, 56}

processed_rows = 0
replaced_cells = 0
cleared_rows = 0

for row_idx, row in enumerate(table.rows):
    # Clear redundant rows
    if row_idx in clear_rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.clear()
        cleared_rows += 1
        continue

    # Placeholder replacement rows
    if row_idx in row_placeholders:
        ph = row_placeholders[row_idx]
        cell_count = len(row.cells)
        print(f"  Row {row_idx}: has {cell_count} cell(s) -> placing {ph}")
        # Replace every cell in this row
        for col_idx, cell in enumerate(row.cells):
            print(f"    DEBUG: col {col_idx} cell id={id(cell)}")
            original = cell.text[:40]
            replace_cell_with_placeholder(cell, ph)
            after = cell.text
            print(f"    col {col_idx}: before='{original!r}' after='{after!r}'")
        if cell_count > 0:
            print(f"    Replaced {cell_count} cells (first: '{original!r}')")
        processed_rows += 1
        replaced_cells += cell_count

print(f"\nSummary: {processed_rows} rows replaced, {replaced_cells} cells total")
print(f"Cleared {cleared_rows} detail rows")

# Save
print(f"\nSaving to: {dst}")
doc.save(dst)
print("Done.")

# === Verification ===
print("\n=== Verification ===")
doc2 = Document(dst)
table2 = doc2.tables[0]
PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
found = {}
for row in table2.rows:
    for cell in row.cells:
        for m in PLACEHOLDER_PATTERN.finditer(cell.text):
            name = m.group(1)
            found[name] = found.get(name, 0) + 1

expected_counts = {
    'SUMMARY': 4,
    'SKILLS_SECTION': 12,   # 3 rows × 4 cols
    'EXP1_BULLET1': 4,
    'EXP1_BULLET2': 4,
    'EXP1_BULLET3': 4,
    'EXP1_BULLET4': 4,
    'EXP1_BULLET5': 4,
    'EXP1_BULLET6': 4,
    'EXP1_BULLET7': 4,
    'EXP1_BULLET8': 4,
    'EXP1_BULLET9': 4,
    'EDUCATION_SECTION': 4,
}
total_expected = sum(expected_counts.values())
total_actual = sum(found.values())

print(f"Placeholder occurrences: expected={total_expected}, actual={total_actual}")
all_ok = True
for name, exp_count in expected_counts.items():
    act = found.get(name, 0)
    match = (act == exp_count)
    symbol = "[OK]" if match else "[MISS]"
    print(f"  {name}: expected {exp_count}, got {act} {symbol}")
    if not match:
        all_ok = False

if all_ok and len(found) == len(expected_counts):
    print("\nAll placeholders inserted correctly with correct column counts [OK]")
else:
    print("\nDiscrepancies found - review output above")
