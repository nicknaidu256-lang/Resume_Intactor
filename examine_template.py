"""
Template Examination Utility
Examines the Master_Resume.docx structure to understand exact formatting and placeholders.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def get_alignment_name(alignment):
    """Convert alignment enum to readable name."""
    if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
        return "LEFT"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        return "CENTER"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        return "RIGHT"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        return "JUSTIFY"
    else:
        return "UNKNOWN"

def examine_paragraph(paragraph, para_idx, context="body"):
    """Examine a single paragraph and its formatting."""
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
    
    if not paragraph.text.strip():
        return None
    
    info = {
        "index": para_idx,
        "context": context,
        "text": paragraph.text,
        "alignment": get_alignment_name(paragraph.alignment),
        "runs": [],
        "placeholders": []
    }
    
    # Examine runs and their formatting
    for run_idx, run in enumerate(paragraph.runs):
        run_info = {
            "index": run_idx,
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name if run.font else None,
            "font_size": run.font.size.pt if run.font and run.font.size else None
        }
        info["runs"].append(run_info)
    
    # Find placeholders
    full_text = paragraph.text
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    for match in matches:
        placeholder_info = {
            "name": match.group(1).strip(),
            "full_match": match.group(0),
            "start": match.start(),
            "end": match.end()
        }
        info["placeholders"].append(placeholder_info)
    
    return info if info["placeholders"] or info["runs"] else None

def examine_table(table, table_idx, level=0):
    """Recursively examine a table and its cells."""
    table_info = {
        "index": table_idx,
        "level": level,
        "rows": len(table.rows),
        "columns": len(table.rows[0].cells) if table.rows else 0,
        "cells": []
    }
    
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_info = {
                "row": row_idx,
                "column": cell_idx,
                "paragraphs": [],
                "nested_tables": []
            }
            
            # Examine cell paragraphs
            for para_idx, paragraph in enumerate(cell.paragraphs):
                para_info = examine_paragraph(paragraph, para_idx, f"table_{table_idx}_cell_{row_idx}_{cell_idx}")
                if para_info:
                    cell_info["paragraphs"].append(para_info)
            
            # Recursively examine nested tables
            for nested_idx, nested_table in enumerate(cell.tables):
                nested_info = examine_table(nested_table, nested_idx, level + 1)
                cell_info["nested_tables"].append(nested_info)
            
            if cell_info["paragraphs"] or cell_info["nested_tables"]:
                table_info["cells"].append(cell_info)
    
    return table_info

def main():
    template_path = Path("templates/Master_Resume.docx")
    
    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return
    
    print(f"Examining template: {template_path}")
    doc = Document(template_path)
    
    # Collect all findings
    findings = {
        "paragraphs": [],
        "tables": [],
        "all_placeholders": set()
    }
    
    # Examine body paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_info = examine_paragraph(paragraph, para_idx)
        if para_info:
            findings["paragraphs"].append(para_info)
            for ph in para_info["placeholders"]:
                findings["all_placeholders"].add(ph["name"])
    
    # Examine tables
    for table_idx, table in enumerate(doc.tables):
        table_info = examine_table(table, table_idx)
        if table_info["cells"]:
            findings["tables"].append(table_info)
            # Extract placeholders from table cells
            for cell in table_info["cells"]:
                for para in cell["paragraphs"]:
                    for ph in para["placeholders"]:
                        findings["all_placeholders"].add(ph["name"])
                for nested in cell["nested_tables"]:
                    for nested_cell in nested["cells"]:
                        for para in nested_cell["paragraphs"]:
                            for ph in para["placeholders"]:
                                findings["all_placeholders"].add(ph["name"])
    
    # Print summary
    print(f"\n=== TEMPLATE ANALYSIS SUMMARY ===")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Unique placeholders found: {len(findings['all_placeholders'])}")
    
    if findings["all_placeholders"]:
        print(f"Placeholders: {sorted(findings['all_placeholders'])}")
    
    # Print detailed info for first few paragraphs with placeholders
    print(f"\n=== DETAILED PARAGRAPH INFO ===")
    placeholder_paras = [p for p in findings["paragraphs"] if p["placeholders"]]
    for para_info in placeholder_paras[:3]:  # Show first 3
        print(f"\nParagraph {para_info['index']} ({para_info['context']}):")
        print(f"  Text: '{para_info['text']}'")
        print(f"  Alignment: {para_info['alignment']}")
        print(f"  Runs: {len(para_info['runs'])}")
        print(f"  Placeholders: {[ph['name'] for ph in para_info['placeholders']]}")
        
        for run in para_info["runs"]:
            fmt = []
            if run["bold"]: fmt.append("bold")
            if run["italic"]: fmt.append("italic") 
            if run["underline"]: fmt.append("underline")
            print(f"    Run {run['index']}: '{run['text']}' {'+'.join(fmt) if fmt else 'normal'}")
    
    # Print table info
    if findings["tables"]:
        print(f"\n=== TABLE INFO ===")
        for table_info in findings["tables"]:
            print(f"Table {table_info['index']}: {table_info['rows']}x{table_info['columns']}, level {table_info['level']}")
            placeholder_cells = [c for c in table_info["cells"] if any(p["placeholders"] for p in c["paragraphs"])]
            if placeholder_cells:
                print(f"  Cells with placeholders: {len(placeholder_cells)}")

if __name__ == "__main__":
    main()