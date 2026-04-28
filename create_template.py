"""
Create a sample Master_Resume.docx template with standard placeholders.
Run this script once to initialize the template.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_master_template():
    """Generate a professional resume template with {{PLACEHOLDER}} sections."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header: Name
    header = doc.add_heading('Abhilash Naidu Paspulati', level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info (static - not tailored)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: abhilash@example.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/abhilash')
    doc.add_paragraph()  # spacer

    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run('{{SUMMARY}}')

    # Experience Section
    doc.add_heading('Professional Experience', level=1)

    # Job 1
    doc.add_heading('Senior Software Engineer', level=2)
    doc.add_paragraph('TechCorp Inc. | 2020 - Present')
    exp1 = doc.add_paragraph()
    exp1.add_run('{{EXP1_BULLET1}}')
    exp2 = doc.add_paragraph()
    exp2.add_run('{{EXP1_BULLET2}}')
    exp3 = doc.add_paragraph()
    exp3.add_run('{{EXP1_BULLET3}}')

    # Job 2
    doc.add_heading('Software Developer', level=2)
    doc.add_paragraph('Innovatech Solutions | 2017 - 2020')
    exp4 = doc.add_paragraph()
    exp4.add_run('{{EXP2_BULLET1}}')
    exp5 = doc.add_paragraph()
    exp5.add_run('{{EXP2_BULLET2}}')
    exp6 = doc.add_paragraph()
    exp6.add_run('{{EXP2_BULLET3}}')

    # Skills Section
    doc.add_heading('Skills', level=1)
    skills_para = doc.add_paragraph()
    skills_para.add_run('{{SKILLS_SECTION}}')

    # Education Section
    doc.add_heading('Education', level=1)
    edu_para = doc.add_paragraph()
    edu_para.add_run('{{EDUCATION_SECTION}}')

    # Save template
    templates_dir = Path('templates')
    templates_dir.mkdir(parents=True, exist_ok=True)
    template_path = templates_dir / 'Master_Resume.docx'
    doc.save(template_path)
    print(f"Template created: {template_path}")

    # Copy to Archive as protected original
    archive_dir = Path('Archive')
    archive_dir.mkdir(parents=True, exist_ok=True)
    archive_path = archive_dir / 'Original_Resume_Master.docx'
    doc.save(archive_path)
    print(f"Protected original saved: {archive_path}")

    return template_path, archive_path


if __name__ == '__main__':
    create_master_template()
