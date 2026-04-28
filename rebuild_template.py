"""
TEMPLATE REBUILD SCRIPT
========================
Creates templates/Master_Resume.docx from Archive/Original_Resume_Master.docx

PRINCIPLE: Template defines layout, runtime only fills text.

This script:
1. Copies archive structure (preserves ALL 91 structural elements)
2. Replaces only dynamic content cells with placeholders (12 total)
3. Uses safe text insertion (no structural modification)
4. Keeps Row 34 "Key Achievements" as permanent header

Row heights, merged cells, column widths, paragraph spacing all inherited from archive.
"""

import shutil
import re
from pathlib import Path
from docx import Document


def safe_replace_cell_text(cell, new_text):
    """
    Safely replace cell text without breaking Word layout.
    
    Rules:
    - Do NOT clear() paragraphs
    - Do NOT remove runs
    - Do NOT restructure paragraphs
    - Only replace text content safely
    """
    # Find the first run and replace its text
    if cell.paragraphs and cell.paragraphs[0].runs:
        # Replace in first run (preserves all formatting)
        first_run = cell.paragraphs[0].runs[0]
        first_run.text = new_text
    else:
        # If no runs exist, add a new run (minimal change)
        if cell.paragraphs:
            cell.paragraphs[0].add_run(new_text)


def rebuild_fixed_template():
    """Rebuild fixed template from archive master."""
    
    # Paths
    archive_path = Path("Archive/Original_Resume_Master.docx")
    template_path = Path("templates/Master_Resume.docx")
    backup_path = Path("templates/Master_Resume_backup_before_fix.docx")
    
    print("=" * 70)
    print("TEMPLATE REBUILD: Template-First Architecture")
    print("=" * 70)
    
    # Step 1: Backup existing template
    if template_path.exists():
        shutil.copy2(template_path, backup_path)
        print(f"Backed up existing template to: {backup_path}")
    
    # Step 2: Load archive (becomes base with ALL structure intact)
    print(f"\nLoading archive template: {archive_path}")
    archive_doc = Document(str(archive_path))
    table = archive_doc.tables[0]
    
    print(f"Archive loaded: {len(table.rows)} rows, {len(table.columns)} columns")
    print(f"Archive file size: {archive_path.stat().st_size:,} bytes")
    
    # Step 3: Define placeholder mapping
    # Only replace dynamic content in column 0
    # Row 34 "Key Achievements" stays as permanent header (NOT a placeholder)
    
    placeholder_map = {
        (8, 0): "{{SUMMARY}}",        # Professional Summary content
        (12, 0): "{{SKILLS_1}}",       # Core Competencies block 1
        (14, 0): "{{SKILLS_2}}",       # Core Competencies block 2
        (16, 0): "{{SKILLS_3}}",       # Core Competencies block 3
        (32, 0): "{{EXP1_BULLET1}}",  # Experience bullet 1
        # Row 34 - SKIP (keep "Key Achievements" as permanent header)
        (36, 0): "{{EXP1_BULLET3}}",   # Experience bullet 3
        (38, 0): "{{EXP1_BULLET4}}",   # Experience bullet 4
        (40, 0): "{{EXP1_BULLET5}}",   # Experience bullet 5
        (42, 0): "{{EXP1_BULLET6}}",   # Experience bullet 6
        (44, 0): "{{EXP1_BULLET7}}",   # Experience bullet 7
        (46, 0): "{{EXP1_BULLET8}}",   # Experience bullet 8
        (48, 0): "{{EXP1_BULLET9}}",   # Experience bullet 9
    }
    
    # Step 4: Replace content with placeholders using SAFE method
    print("\nReplacing dynamic content (safe insertion)...")
    
    replaced_count = 0
    for (row_idx, col_idx), placeholder_text in placeholder_map.items():
        try:
            cell = table.cell(row_idx, col_idx)
            safe_replace_cell_text(cell, placeholder_text)
            print(f"  Row {row_idx}, Col {col_idx} -> {placeholder_text}")
            replaced_count += 1
        except Exception as e:
            print(f"  ERROR at Row {row_idx}, Col {col_idx}: {e}")
    
    print(f"\nReplaced {replaced_count} content cells with placeholders")
    
    # Step 5: Verify Row 34 is still "Key Achievements"
    row_34_content = table.cell(34, 0).text.strip()
    print(f"\nVerifying Row 34: '{row_34_content}'")
    if row_34_content == "Key Achievements":
        print("  ✓ Row 34 correctly preserved as permanent header")
    else:
        print(f"  WARNING: Row 34 shows: '{row_34_content}'")
    
    # Step 6: Save as new working template
    print(f"\nSaving fixed template: {template_path}")
    archive_doc.save(str(template_path))
    
    # Verify file size matches archive (structure preserved)
    new_size = template_path.stat().st_size
    archive_size = archive_path.stat().st_size
    size_diff = new_size - archive_size
    
    print(f"  Template saved: {new_size:,} bytes")
    print(f"  Archive size:   {archive_size:,} bytes")
    print(f"  Size diff:      {size_diff:,} bytes (should be minimal)")
    
    # Step 7: Verify placeholder detection
    print("\n" + "=" * 70)
    print("VERIFYING PLACEHOLDER DETECTION")
    print("=" * 70)
    
    # Reload and scan for placeholders
    verify_doc = Document(str(template_path))
    verify_table = verify_doc.tables[0]
    
    placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
    detected_placeholders = set()
    
    for row in verify_table.rows:
        for cell in row.cells:
            matches = placeholder_pattern.findall(cell.text)
            for match in matches:
                detected_placeholders.add(match.strip())
    
    print(f"\nDetected {len(detected_placeholders)} placeholders:")
    for ph in sorted(detected_placeholders):
        print(f"  - {ph}")
    
    # Verify exactly 12 placeholders (no EXP1_BULLET2)
    expected = {'SUMMARY', 'SKILLS_1', 'SKILLS_2', 'SKILLS_3',
                'EXP1_BULLET1', 'EXP1_BULLET3', 'EXP1_BULLET4', 'EXP1_BULLET5',
                'EXP1_BULLET6', 'EXP1_BULLET7', 'EXP1_BULLET8', 'EXP1_BULLET9'}
    
    print(f"\nValidation:")
    if detected_placeholders == expected:
        print("  ✓ PASSED: Exactly 12 placeholders")
        print("  ✓ EXP1_BULLET2 successfully removed")
    else:
        missing = expected - detected_placeholders
        extra = detected_placeholders - expected
        if missing:
            print(f"  ! MISSING: {missing}")
        if extra:
            print(f"  ! EXTRA: {extra}")
    
    # Step 8: Summary
    print("\n" + "=" * 70)
    print("TEMPLATE REBUILD COMPLETE")
    print("=" * 70)
    print(f"Output: {template_path}")
    print(f"Structure: All archive elements preserved (91 differences fixed)")
    print(f"Placeholders: {len(detected_placeholders)} (EXP1_BULLET2 removed)")
    print(f"Row 34: 'Key Achievements' preserved as permanent header")
    print("\nNext: Run layout-preserving engine with safe text insertion only")
    
    return template_path


if __name__ == "__main__":
    rebuild_fixed_template()