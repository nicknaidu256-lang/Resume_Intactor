"""
TEMPLATE REBUILD SCRIPT - FIXED VERSION
========================================
Creates templates/Master_Resume.docx from Archive/Original_Resume_Master.docx

PRINCIPLE: Template defines layout, runtime only fills text.

Key fix: The previous script used wrong row numbers. This version uses the
ACTUAL positions found in the template.
"""

import shutil
import re
from pathlib import Path
from docx import Document


def safe_replace_cell_text(cell, new_text):
    """Safely replace cell text without breaking Word layout."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        first_run = cell.paragraphs[0].runs[0]
        first_run.text = new_text
    else:
        if cell.paragraphs:
            cell.paragraphs[0].add_run(new_text)


def rebuild_fixed_template():
    """Rebuild fixed template from archive master with CORRECT positions."""
    
    archive_path = Path("Archive/Original_Resume_Master.docx")
    template_path = Path("templates/Master_Resume.docx")
    backup_path = Path("templates/Master_Resume_v2_backup.docx")
    
    print("=" * 70)
    print("TEMPLATE REBUILD - FIXED POSITIONS")
    print("=" * 70)
    
    # Backup
    if template_path.exists():
        shutil.copy2(template_path, backup_path)
        print(f"Backed up to: {backup_path}")
    
    # Load archive
    archive_doc = Document(str(archive_path))
    table = archive_doc.tables[0]
    
    print(f"Archive: {len(table.rows)} rows, {len(table.columns)} columns")
    
    # CORRECT placeholder mapping based on actual positions
    # Only replace dynamic content - keep section headers as-is
    # Row 34 "Key Achievements" stays as permanent header
    
    placeholder_map = {
        # Summary section - Row 8, Col 0 (dynamic content)
        (8, 0): "{{SUMMARY}}",
        
        # Skills section - spread across cols 1, 2, 3
        (11, 1): "{{SKILLS_1}}",
        (12, 2): "{{SKILLS_2}}",
        (13, 3): "{{SKILLS_3}}",
        
        # Experience bullets - spread across rows and cols
        (29, 0): "{{EXP1_BULLET1}}",
        # Row 30, Col 1 - this is job title, NOT dynamic - SKIP
        (31, 2): "{{EXP1_BULLET3}}",
        (32, 3): "{{EXP1_BULLET4}}",
        # Row 34, Col 0 - "Key Achievements" header - SKIP (keep as permanent)
        (35, 1): "{{EXP1_BULLET6}}",
        # Row 36, Col 2 - this is section content - SKIP (keep archive)
        (37, 3): "{{EXP1_BULLET8}}",
        (39, 0): "{{EXP1_BULLET9}}",
    }
    
    # Additional dynamic placeholders to add
    # The original template had EXP1_BULLET5, EXP1_BULLET7 as placeholders
    # but archive has content there - we need to check what's truly dynamic
    
    # Map: Where to put placeholder -> archive content -> should be dynamic?
    dynamic_replacements = []
    
    # Check each position in archive vs template
    print("\nAnalyzing content...")
    
    # SUMMARY
    dynamic_replacements.append(((8, 0), "{{SUMMARY}}"))
    
    # SKILLS - in cols 1, 2, 3 at rows 11, 12, 13
    dynamic_replacements.append(((11, 1), "{{SKILLS_1}}"))
    dynamic_replacements.append(((12, 2), "{{SKILLS_2}}"))
    dynamic_replacements.append(((13, 3), "{{SKILLS_3}}"))
    
    # EXPERIENCE - this is complex, need to be careful
    # Row 29-32 have bullets in archive, need to identify which are placeholders
    
    print("\nReplacing dynamic content...")
    replaced = 0
    for (row, col), placeholder in dynamic_replacements:
        try:
            cell = table.cell(row, col)
            safe_replace_cell_text(cell, placeholder)
            print(f"  Row {row}, Col {col} -> {placeholder}")
            replaced += 1
        except Exception as e:
            print(f"  ERROR Row {row}, Col {col}: {e}")
    
    print(f"\nReplaced {replaced} cells")
    
    # Verify Row 34 - should be "Key Achievements"
    row_34 = table.cell(34, 0).text.strip()
    print(f"\nRow 34 check: '{row_34}'")
    if row_34 == "Key Achievements":
        print("  OK - Preserved as permanent header")
    else:
        print(f"  WARNING: Expected 'Key Achievements', got '{row_34}'")
    
    # Save
    print(f"\nSaving: {template_path}")
    archive_doc.save(str(template_path))
    
    # Verify
    print("\n" + "=" * 70)
    print("VERIFYING")
    print("=" * 70)
    
    verify_doc = Document(str(template_path))
    v_table = verify_doc.tables[0]
    
    pattern = re.compile(r'\{\{([^}]+)\}\}')
    found = set()
    
    for row in v_table.rows:
        for cell in row.cells:
            for m in pattern.findall(cell.text):
                found.add(m.strip())
    
    print(f"Found {len(found)} placeholders:")
    for p in sorted(found):
        print(f"  - {p}")
    
    # Check Row 34 specifically
    row34 = v_table.cell(34, 0).text.strip()
    print(f"\nRow 34 in final: '{row34}'")
    
    return template_path


if __name__ == "__main__":
    rebuild_fixed_template()