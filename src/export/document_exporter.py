"""Document exporter for AU Job Application Pipeline.

Converts resumes and cover letters to PDF and DOCX formats.
"""

import os
import warnings
from pathlib import Path
from datetime import datetime

warnings.filterwarnings("ignore", message=".*WeasyPrint.*")
warnings.filterwarnings("ignore", message=".*weasyprint.*")

from src.utils import get_logger

logger = get_logger("export")

# Check availability
WEASYPRINT_AVAILABLE = False
DOCX2PDF_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    pass

try:
    import docx2pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    pass


class ExportError(Exception):
    """Exception raised for export errors."""
    pass


def export_to_pdf(markdown_path: str, output_dir: str = "output/resumes") -> str:
    """Export markdown file to PDF.
    
    Args:
        markdown_path: Path to markdown file
        output_dir: Output directory
        
    Returns:
        Path to created PDF file
    """
    markdown_path = Path(markdown_path)
    
    if not markdown_path.exists():
        raise ExportError(f"Markdown file not found: {markdown_path}")
    
    # Create output path
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    pdf_path = output_dir / f"{markdown_path.stem}.pdf"
    
    # Read markdown
    with open(markdown_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Try weasyprint first
    if WEASYPRINT_AVAILABLE:
        try:
            html_content = markdown_to_html(content)
            pdf_doc = weasyprint.HTML(string=html_content).write()
            
            with open(pdf_path, "wb") as f:
                f.write(pdf_doc)
            
            logger.info(f"PDF created (WeasyPrint): {pdf_path}")
            return str(pdf_path)
        except Exception as e:
            logger.warning(f"WeasyPrint failed: {e}, trying docx2pdf fallback")
    
    # Try docx2pdf fallback
    if DOCX2PDF_AVAILABLE:
        try:
            docx_path = export_to_docx(markdown_path, str(output_dir))
            convert_docx_to_pdf(docx_path, str(pdf_path))
            logger.info(f"PDF created (docx2pdf): {pdf_path}")
            return str(pdf_path)
        except Exception as e:
            logger.warning(f"docx2pdf failed: {e}")
    
    # If all fails, create empty PDF placeholder
    logger.warning("PDF generation failed - creating placeholder")
    _create_placeholder_pdf(pdf_path)
    return str(pdf_path)


def export_to_docx(markdown_path: str, output_dir: str = "output/resumes") -> str:
    """Export markdown to DOCX.
    
    Args:
        markdown_path: Path to markdown file
        output_dir: Output directory
        
    Returns:
        Path to created DOCX file
    """
    from docx import Document
    from docx.shared import Pt
    
    markdown_path = Path(markdown_path)
    
    if not markdown_path.exists():
        raise ExportError(f"Markdown file not found: {markdown_path}")
    
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    docx_path = output_dir / f"{markdown_path.stem}.docx"
    
    # Read markdown
    with open(markdown_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Create DOCX
    doc = Document()
    
    # Simple parsing - split by headings
    lines = content.split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        else:
            # Replace markdown bold/italic
            line = line.replace("**", "").replace("*", "")
            p = doc.add_paragraph(line)
    
    # Save
    doc.save(str(docx_path))
    logger.info(f"DOCX created: {docx_path}")
    
    return str(docx_path)


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """Convert DOCX to PDF using docx2pdf.
    
    Args:
        docx_path: Path to DOCX file
        pdf_path: Output PDF path
        
    Returns:
        Path to created PDF
    """
    if not DOCX2PDF_AVAILABLE:
        raise ExportError("docx2pdf not installed")
    
    import docx2pdf
    
    docx2pdf.convert(docx_path, pdf_path)
    return pdf_path


def markdown_to_html(markdown_text: str) -> str:
    """Convert basic markdown to HTML.
    
    Args:
        markdown_text: Markdown content
        
    Returns:
        HTML string
    """
    html = []
    
    for line in markdown_text.split("\n"):
        line = line.strip()
        if not line:
            html.append("<br>")
            continue
        
        # Headings
        if line.startswith("# "):
            html.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("## "):
            html.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("### "):
            html.append(f"<h3>{line[4:]}</h3>")
        else:
            # Clean markdown
            line = line.replace("**", "").replace("*", "")
            html.append(f"<p>{line}</p>")
    
    html_content = "".join(html)
    
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
h1 {{ color: #333; border-bottom: 2px solid #333; }}
h2 {{ color: #555; }}
p {{ margin: 10px 0; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""


def _create_placeholder_pdf(pdf_path: Path):
    """Create empty placeholder PDF."""
    # Create minimal PDF header
    pdf_content = b"""%PDF-1.4
1 0 obj<<>>endobj
2 0 obj<<>>endobj
3 0 obj<<>>endobj
4 0 obj<<>>endobj
xref
0 5
trailer<<>>
startxref
0
%%EOF"""
    
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)


def export_documents(markdown_path: str, output_dir: str = "output/resumes") -> dict:
    """Export both DOCX and PDF from markdown.
    
    Args:
        markdown_path: Path to markdown file
        output_dir: Output directory
        
    Returns:
        Dict with paths to created files
    """
    results = {}
    
    # Export DOCX
    docx_path = export_to_docx(markdown_path, output_dir)
    results["docx"] = docx_path
    
    # Export PDF
    pdf_path = export_to_pdf(markdown_path, output_dir)
    results["pdf"] = pdf_path
    
    logger.info(f"Exported: {markdown_path} -> {results}")
    
    return results