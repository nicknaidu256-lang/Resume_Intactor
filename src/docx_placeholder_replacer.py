"""
docx_placeholder_replacer.py

Run-safe placeholder replacement for Word (.docx) templates using python-docx.

Key constraint: never use paragraph.text = ... for final replacements because it
collapses runs and can destroy formatting. Instead we:
  - find placeholders in the concatenated run text
  - split boundary runs at exact character offsets
  - replace the placeholder span by editing/removing only the affected runs

This preserves all formatting, spacing, alignment, table layout, and header/footer
structure outside the placeholder span.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from typing import Dict, Iterable, Iterator, List, Optional, Tuple

from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# {{PLACEHOLDER_NAME}}
PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")


@dataclass(frozen=True)
class PlaceholderMatch:
    name: str
    start: int  # char offset in paragraph concatenated run text
    end: int    # char offset in paragraph concatenated run text (exclusive)


def iter_all_paragraphs(doc: _Document) -> Iterator[Paragraph]:
    """
    Yield every Paragraph in a document:
    - body paragraphs
    - table cell paragraphs (including nested tables)
    - headers and footers for all sections
    """
    yield from doc.paragraphs

    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    for section in doc.sections:
        header = section.header
        footer = section.footer
        yield from header.paragraphs
        for table in header.tables:
            yield from _iter_table_paragraphs(table)
        yield from footer.paragraphs
        for table in footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterator[Paragraph]:
    yield from cell.paragraphs
    for nested in cell.tables:
        yield from _iter_table_paragraphs(nested)


def scan_placeholders(doc: _Document) -> Dict[str, int]:
    """
    Scan the entire document and return counts per placeholder name.
    """
    counts: Dict[str, int] = {}
    for p in iter_all_paragraphs(doc):
        text = _paragraph_full_text(p)
        if not text:
            continue
        for m in PLACEHOLDER_PATTERN.finditer(text):
            name = m.group(1).strip()
            counts[name] = counts.get(name, 0) + 1
    return counts


def replace_placeholders_in_document(doc: _Document, replacements: Dict[str, str]) -> Dict[str, int]:
    """
    Replace placeholders across the entire document (body + tables + headers/footers).

    Returns stats:
      - replaced: number of placeholder occurrences replaced
      - skipped_missing: occurrences with no replacement provided (left unchanged)
      - skipped_mismatch: occurrences whose matched slice didn't look like {{...}}
    """
    stats = {"replaced": 0, "skipped_missing": 0, "skipped_mismatch": 0}
    for paragraph in iter_all_paragraphs(doc):
        r = replace_placeholders_in_paragraph(paragraph, replacements)
        stats["replaced"] += r["replaced"]
        stats["skipped_missing"] += r["skipped_missing"]
        stats["skipped_mismatch"] += r["skipped_mismatch"]
    return stats


def replace_literal_in_document(doc: _Document, literal: str, replacement: str) -> int:
    """
    Run-safe literal string replacement across the whole document.

    This is used as a backward-compatible fallback for templates that still contain
    fixed text (e.g., a headline) instead of a {{PLACEHOLDER}}.

    Returns the number of literal occurrences replaced.
    """
    if not literal:
        return 0
    count = 0
    for paragraph in iter_all_paragraphs(doc):
        count += replace_literal_in_paragraph(paragraph, literal, replacement)
    return count


def replace_literal_in_paragraph(paragraph: Paragraph, literal: str, replacement: str) -> int:
    """
    Replace all occurrences of a literal substring within a paragraph (run-safe).
    """
    if not paragraph.runs:
        return 0

    replaced = 0
    full_text = _paragraph_full_text(paragraph)
    if not full_text or literal not in full_text:
        return 0

    # Replace right-to-left to keep offsets stable.
    start = full_text.rfind(literal)
    while start != -1:
        end = start + len(literal)
        _replace_range_across_runs(paragraph, start, end, replacement)
        replaced += 1
        full_text = full_text[:start] + replacement + full_text[end:]
        start = full_text.rfind(literal, 0, start)

    return replaced


def replace_placeholders_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> Dict[str, int]:
    """
    Replace placeholders in a single paragraph by editing runs safely.
    """
    stats = {"replaced": 0, "skipped_missing": 0, "skipped_mismatch": 0}
    if not paragraph.runs:
        return stats

    full_text = _paragraph_full_text(paragraph)
    if not full_text:
        return stats

    matches: List[PlaceholderMatch] = []
    for m in PLACEHOLDER_PATTERN.finditer(full_text):
        name = m.group(1).strip()
        matches.append(PlaceholderMatch(name=name, start=m.start(), end=m.end()))

    if not matches:
        return stats

    # Replace right-to-left to keep offsets stable.
    for match in reversed(matches):
        if match.name not in replacements:
            stats["skipped_missing"] += 1
            continue
        replacement_text = "" if replacements[match.name] is None else str(replacements[match.name])
        if not PLACEHOLDER_PATTERN.fullmatch(full_text[match.start : match.end]):
            stats["skipped_mismatch"] += 1
            continue

        _replace_range_across_runs(paragraph, match.start, match.end, replacement_text)
        stats["replaced"] += 1

        # Update cached full_text for subsequent (earlier) replacements in same paragraph.
        full_text = full_text[: match.start] + replacement_text + full_text[match.end :]

    return stats


def _paragraph_full_text(paragraph: Paragraph) -> str:
    # Avoid paragraph.text because python-docx may normalize; we want run-exact concatenation.
    return "".join(run.text for run in paragraph.runs)


def _run_position(paragraph: Paragraph, char_index: int, *, prefer_next_boundary: bool) -> Tuple[int, int]:
    """
    Map a character index in the concatenated run text to (run_index, offset_in_run).

    If char_index lands exactly at a run boundary:
      - prefer_next_boundary=True  -> point at the start of the next run
      - prefer_next_boundary=False -> point at the end of the previous run
    """
    if char_index < 0:
        raise ValueError("char_index must be >= 0")

    runs = paragraph.runs
    cumulative = 0
    last_non_empty = None

    for i, run in enumerate(runs):
        text = run.text or ""
        run_len = len(text)
        if run_len == 0:
            continue
        last_non_empty = i

        start = cumulative
        end = cumulative + run_len

        if start <= char_index < end:
            return i, char_index - start

        if char_index == end:
            if prefer_next_boundary:
                # Prefer the next non-empty run, if any.
                nxt = _next_non_empty_run_index(runs, i + 1)
                if nxt is not None:
                    return nxt, 0
                # Boundary at end of paragraph: map to this run's end.
                return i, run_len
            return i, run_len

        cumulative = end

    # If index points at the very end, map to last non-empty run end.
    if last_non_empty is not None and char_index == cumulative:
        return last_non_empty, len(runs[last_non_empty].text or "")

    raise ValueError(f"char_index {char_index} out of range for paragraph length {cumulative}")


def _next_non_empty_run_index(runs: List[Run], start: int) -> Optional[int]:
    for j in range(start, len(runs)):
        if runs[j].text:
            return j
    return None


def _split_run(run: Run, split_at: int) -> Run:
    """
    Split a run into two runs at split_at (0..len(run.text)).
    Returns the newly created run (the right side) inserted immediately after `run`.
    """
    text = run.text or ""
    if split_at < 0 or split_at > len(text):
        raise ValueError("split_at out of range")

    if split_at == len(text):
        # Nothing on the right.
        return run
    if split_at == 0:
        # Everything on the right; create a clone before modifying text.
        new_r = deepcopy(run._r)
        run._r.addnext(new_r)
        new_run = Run(new_r, run._parent)
        new_run.text = text
        run.text = ""
        return new_run

    left = text[:split_at]
    right = text[split_at:]

    new_r = deepcopy(run._r)
    run._r.addnext(new_r)
    new_run = Run(new_r, run._parent)

    run.text = left
    new_run.text = right
    return new_run


def _remove_run(run: Run) -> None:
    parent = run._r.getparent()
    if parent is not None:
        parent.remove(run._r)

def _run_index_by_element(paragraph: Paragraph, run_element) -> int:
    """
    Find the index of a run in paragraph.runs by comparing the underlying XML element.
    This is more reliable than list.index(Run) because python-docx may wrap the same
    <w:r> element with different Run objects across accesses.
    """
    for i, r in enumerate(paragraph.runs):
        if r._r is run_element:
            return i
    raise ValueError("run element not found in paragraph")


def _replace_range_across_runs(paragraph: Paragraph, start: int, end: int, replacement: str) -> None:
    """
    Replace character range [start, end) in concatenated run text with `replacement`,
    by splitting boundary runs and removing/altering only runs that intersect the range.
    """
    if start >= end:
        return

    # Resolve positions using current paragraph runs.
    # - start: if it lands on a run boundary, prefer the next run (start of placeholder)
    # - end: if it lands on a run boundary, prefer the previous run (end of placeholder)
    start_run_idx, start_off = _run_position(paragraph, start, prefer_next_boundary=True)
    end_run_idx, end_off = _run_position(paragraph, end, prefer_next_boundary=False)

    # Refresh run references after any splits.
    runs = paragraph.runs
    start_run = runs[start_run_idx]
    end_run = runs[end_run_idx]
    start_el = start_run._r
    end_el = end_run._r

    # Split end run first so indices for start remain valid if start/end in same run.
    tail_run: Optional[Run] = None
    if end_off < len(end_run.text or ""):
        tail_run = _split_run(end_run, end_off)
        # After splitting, tail_run is the right part (kept), end_run becomes left.
        # Placeholder span should not include tail_run.
        runs = paragraph.runs
        # If end_run was before start_run due to boundary mapping, recompute.
        start_run_idx, start_off = _run_position(paragraph, start, prefer_next_boundary=True)
        end_run_idx, _ = _run_position(paragraph, end, prefer_next_boundary=False)
        start_run = paragraph.runs[start_run_idx]
        end_run = paragraph.runs[end_run_idx]
        start_el = start_run._r
        end_el = end_run._r

    if start_off > 0:
        start_run = _split_run(start_run, start_off)
        runs = paragraph.runs
        start_run_idx = _run_index_by_element(paragraph, start_run._r)
        # end index may have shifted by +1 if split occurred before it
        end_run_idx, _ = _run_position(paragraph, end, prefer_next_boundary=False)
        end_run = paragraph.runs[end_run_idx]
        start_el = start_run._r
        end_el = end_run._r

    # Now the placeholder starts at the beginning of start_run and ends at the end of end_run.
    start_run.text = replacement

    # Remove runs strictly between start_run and end_run (inclusive of end_run if different).
    runs = paragraph.runs
    start_i = _run_index_by_element(paragraph, start_el)
    end_i = _run_index_by_element(paragraph, end_el)

    for i in range(end_i, start_i, -1):
        _remove_run(paragraph.runs[i])
