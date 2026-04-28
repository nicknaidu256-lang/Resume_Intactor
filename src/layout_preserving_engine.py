"""
LAYOUT-PRESERVING DOCX ENGINE — Exact formatting preservation for resume templates

Specialized engine that maintains EXACT layout, spacing, and formatting 
from Archive/Original_Resume_Master.docx while replacing placeholders.
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from collections import defaultdict
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
# from docx.oxml.ns import qn
# from docx.oxml.xmlchemy import OxmlElement
from docx.shared import RGBColor, Pt

from src.utils import get_logger, ensure_dir_exists

logger = get_logger()

class LayoutPreservingEngine:
    """
    Specialized engine that preserves EXACT layout from archive template.
    
    Key features:
    - Maintains identical row heights, cell widths, margins, and spacing
    - Preserves all formatting: fonts, bold, colors, alignment  
    - Prevents layout collapse when inserting longer content
    - Uses archive template as layout reference
    """
    
    def __init__(self, template_path: Path, archive_path: Path):
        self.template_path = Path(template_path)
        self.archive_path = Path(archive_path)
        self.document = None
        self.archive_document = None
        self.placeholders: Dict[str, List[dict]] = {}
        self.placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
        
        self._load_documents()
        self._scan_placeholders()
    
    def _load_documents(self) -> None:
        """Load both working template and archive reference."""
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
        if not self.archive_path.exists():
            raise FileNotFoundError(f"Archive template not found: {self.archive_path}")
        
        logger.info(f"Loading working template: {self.template_path}")
        self.document = Document(str(self.template_path))
        
        logger.info(f"Loading archive reference: {self.archive_path}")
        self.archive_document = Document(str(self.archive_path))
        
        # Verify both documents have the same structure
        self._validate_structure()
    
    def _validate_structure(self) -> None:
        """Verify that both documents have identical table structure."""
        if len(self.document.tables) != len(self.archive_document.tables):
            raise ValueError("Template and archive have different table counts")
        
        template_table = self.document.tables[0]
        archive_table = self.archive_document.tables[0]
        
        if len(template_table.rows) != len(archive_table.rows):
            raise ValueError("Template and archive have different row counts")
        if len(template_table.columns) != len(archive_table.columns):
            raise ValueError("Template and archive have different column counts")
        
        logger.info(f"Structure validated: {len(template_table.rows)}x{len(template_table.columns)} table")
    
    def _scan_placeholders(self) -> None:
        """Scan for placeholders in the working template."""
        logger.info("Scanning for placeholders...")
        
        table = self.document.tables[0]
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if not paragraph.text.strip():
                        continue
                    
                    matches = list(self.placeholder_pattern.finditer(paragraph.text))
                    for match in matches:
                        placeholder_name = match.group(1).strip()
                        start_pos, end_pos = match.span()
                        
                        occurrence = {
                            "table": table,
                            "row_idx": row_idx,
                            "col_idx": col_idx,
                            "para_idx": para_idx,
                            "paragraph": paragraph,
                            "start_pos": start_pos,
                            "end_pos": end_pos,
                            "full_text": match.group(0)
                        }
                        
                        if placeholder_name not in self.placeholders:
                            self.placeholders[placeholder_name] = []
                        self.placeholders[placeholder_name].append(occurrence)
                        
                        logger.debug(f"Found '{placeholder_name}' at [{row_idx},{col_idx}] position {start_pos}-{end_pos}")
        
        unique_placeholders = sorted(self.placeholders.keys())
        total_occurrences = sum(len(occs) for occs in self.placeholders.values())
        
        logger.info(f"Found {len(unique_placeholders)} unique placeholders with {total_occurrences} total occurrences")
        logger.info(f"Placeholders: {unique_placeholders}")
    
    def replace_placeholders(self, replacements: Dict[str, str]) -> Dict[str, int]:
        """
        Replace placeholders while preserving exact layout.
        
        Strategy:
        1. Use archive template as layout reference
        2. Preserve row heights and cell widths
        3. Maintain identical formatting and spacing
        4. Handle text expansion without collapsing layout
        """
        logger.info("Starting layout-preserving replacement...")
        
        stats = {
            "total_placeholders": sum(len(occs) for occs in self.placeholders.values()),
            "replaced": 0,
            "skipped_missing": 0,
            "skipped_errors": 0
        }
        
        # First, ensure layout matches archive
        self._sync_layout_from_archive()
        
        # Process each placeholder
        for placeholder_name, occurrences in self.placeholders.items():
            replacement_text = replacements.get(placeholder_name)
            
            if replacement_text is None:
                stats["skipped_missing"] += len(occurrences)
                logger.warning(f"No replacement for '{placeholder_name}' — keeping original")
                continue
            
            for occurrence in occurrences:
                try:
                    success = self._replace_single_occurrence(occurrence, replacement_text)
                    if success:
                        stats["replaced"] += 1
                        logger.debug(f"Replaced '{placeholder_name}' successfully")
                    else:
                        stats["skipped_errors"] += 1
                        logger.warning(f"Failed to replace '{placeholder_name}'")
                except Exception as e:
                    stats["skipped_errors"] += 1
                    logger.error(f"Error replacing '{placeholder_name}': {e}")
        
        logger.info(f"Replacement complete: {stats['replaced']} replaced, "
                    f"{stats['skipped_missing']} skipped (missing), "
                    f"{stats['skipped_errors']} skipped (errors)")
        
        return stats
    
    def _sync_layout_from_archive(self) -> None:
        """Sync layout properties from archive template."""
        template_table = self.document.tables[0]
        archive_table = self.archive_document.tables[0]
        
        # Sync row heights
        for row_idx, (template_row, archive_row) in enumerate(zip(template_table.rows, archive_table.rows)):
            if hasattr(archive_row, 'height') and archive_row.height:
                template_row.height = archive_row.height
            
            # Sync cell properties
            for col_idx, (template_cell, archive_cell) in enumerate(zip(template_row.cells, archive_row.cells)):
                # Sync cell width
                if hasattr(archive_cell, 'width') and archive_cell.width:
                    template_cell.width = archive_cell.width
                
                # Sync vertical alignment
                if hasattr(archive_cell, 'vertical_alignment'):
                    template_cell.vertical_alignment = archive_cell.vertical_alignment
        
        logger.debug("Layout synchronized from archive template")
    
    def _detect_duplication_pattern(self, row_idx: int) -> Dict[int, List[int]]:
        """
        Detect content duplication patterns across columns in a row.
        Returns: {content_hash: [column_indices]} mapping
        """
        patterns = defaultdict(list)
        table = self.archive_document.tables[0]
        
        for col_idx in range(len(table.columns)):
            try:
                cell = table.cell(row_idx, col_idx)
                content = cell.text.strip()
                if content:  # Only consider cells with content
                    content_hash = hash(content)
                    patterns[content_hash].append(col_idx)
            except (IndexError, AttributeError):
                continue
        
        return patterns
    
    def _replace_in_cell(self, row_idx: int, col_idx: int, replacement_text: str) -> bool:
        """Replace content in a specific cell while preserving formatting."""
        try:
            template_cell = self.document.tables[0].cell(row_idx, col_idx)
            archive_cell = self.archive_document.tables[0].cell(row_idx, col_idx)
            
            # Clear existing content
            if template_cell.paragraphs:
                for paragraph in template_cell.paragraphs:
                    self._clear_paragraph_runs(paragraph)
            
            # Add new content
            if template_cell.paragraphs:
                # Use first paragraph
                new_run = template_cell.paragraphs[0].add_run(replacement_text)
            else:
                # Create new paragraph
                paragraph = template_cell.add_paragraph()
                new_run = paragraph.add_run(replacement_text)
            
            # Apply formatting from archive
            if archive_cell.paragraphs:
                archive_para = archive_cell.paragraphs[0]
                if archive_para.runs:
                    self._copy_run_formatting(new_run, archive_para.runs[0])
            
            return True
            
        except Exception as e:
            logger.error(f"Error replacing cell [{row_idx},{col_idx}]: {e}")
            return False
    
    def _replace_single_occurrence(self, occurrence: dict, replacement_text: str) -> bool:
        """Replace a single placeholder with layout preservation, handling duplication."""
        row_idx = occurrence["row_idx"]
        col_idx = occurrence["col_idx"]
        
        # Detect duplication pattern from archive
        patterns = self._detect_duplication_pattern(row_idx)
        
        # Find which pattern this cell belongs to
        try:
            archive_cell = self.archive_document.tables[0].cell(row_idx, col_idx)
            original_content = archive_cell.text.strip()
            original_hash = hash(original_content)
            
            if original_hash in patterns and len(patterns[original_hash]) > 1:
                # This was part of a duplication pattern - replace in all columns
                duplicate_columns = patterns[original_hash]
                logger.debug(f"Found duplication pattern in row {row_idx}: replacing in columns {duplicate_columns}")
                
                success_count = 0
                for dup_col in duplicate_columns:
                    if self._replace_in_cell(row_idx, dup_col, replacement_text):
                        success_count += 1
                
                return success_count > 0
            else:
                # Single column - normal replacement
                return self._replace_in_cell(row_idx, col_idx, replacement_text)
                
        except Exception as e:
            logger.error(f"Error in duplication detection for row {row_idx}, col {col_idx}: {e}")
            # Fall back to single cell replacement
            return self._replace_in_cell(row_idx, col_idx, replacement_text)
    
    def _clear_paragraph_runs(self, paragraph: Paragraph) -> None:
        """Clear all runs from a paragraph while preserving paragraph formatting."""
        p_element = paragraph._p
        # Remove all run elements but keep paragraph properties
        for run in paragraph.runs:
            r_element = run._r
            p_element.remove(r_element)
    
    def _copy_run_formatting(self, target_run: Run, source_run: Run) -> None:
        """Copy formatting from source run to target run."""
        if not source_run.font:
            return
        
        target_font = target_run.font
        source_font = source_run.font
        
        # Copy basic formatting
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        
        # Copy font properties
        if source_font.name:
            target_font.name = source_font.name
        if source_font.size:
            target_font.size = source_font.size
        if hasattr(source_font.color, 'rgb') and source_font.color.rgb:
            target_font.color.rgb = source_font.color.rgb
    
    def save(self, output_path: Path) -> None:
        """Save the modified document with preserved layout."""
        if not self.document:
            raise ValueError("No document loaded")
        
        logger.info(f"Saving document with preserved layout: {output_path}")
        ensure_dir_exists(output_path.parent)
        self.document.save(str(output_path))
        logger.info(f"Document saved successfully: {output_path}")
    
    def get_placeholder_names(self) -> List[str]:
        """Get list of all detected placeholder names."""
        return sorted(self.placeholders.keys())
    
    def get_placeholder_count(self, name: str) -> int:
        """Get count of occurrences for a specific placeholder."""
        return len(self.placeholders.get(name, []))


def main():
    """Example usage."""
    template_path = Path("templates/Master_Resume.docx")
    archive_path = Path("Archive/Original_Resume_Master.docx")
    
    try:
        engine = LayoutPreservingEngine(template_path, archive_path)
        
        # Example replacements that match the archive content style
        replacements = {
            "SUMMARY": "Accomplished Engineering Manager with 10 years of experience in pharmaceutical "
                      "and manufacturing industries. Expertise in systems integration, technical governance, "
                      "and capital project delivery. Proven track record in leading cross-functional teams "
                      "and ensuring regulatory compliance.",
            "EXP1_BULLET1": "Lead technology transfer and systems integration for $800M capital program, "
                          "ensuring seamless implementation of new manufacturing capabilities",
            "EXP1_BULLET2": "Managed multi-stream technical delivery across engineering, validation, "
                          "and quality assurance teams to meet project milestones",
            "SKILLS_1": "Engineering Management: Systems Integration, Technical Governance, Capital Project "
                      "Delivery, Team Leadership, Vendor Management",
            "SKILLS_2": "Technical Expertise: Process Engineering, Equipment Qualification, Regulatory "
                      "Compliance, Quality Systems, Risk Management",
            "SKILLS_3": "Professional Skills: Cross-Functional Collaboration, Stakeholder Engagement, "
                      "Problem Solving, Continuous Improvement, Documentation Excellence"
        }
        
        stats = engine.replace_placeholders(replacements)
        print(f"Replacement stats: {stats}")
        
        output_path = Path("output/Layout_Preserved_Resume.docx")
        engine.save(output_path)
        print(f"Saved to: {output_path}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()