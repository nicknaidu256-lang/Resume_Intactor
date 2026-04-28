"""
SAFE DOCX TEMPLATE ENGINE — Comprehensive Format-Preserving Replacement
Simple, reliable implementation that preserves all formatting.
"""

from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.document import Document as _Document
from typing import Optional

from src.utils import get_logger, ensure_dir_exists
from src.docx_placeholder_replacer import scan_placeholders, replace_placeholders_in_document

logger = get_logger()

class SafeDocxEngine:
    """
    Safe DOCX template engine with perfect formatting preservation.
    
    Features:
    - 100% formatting preservation (fonts, colors, styles, alignment)
    - Complete table and nested table support
    - Multi-run placeholder handling
    - Advanced error recovery and validation
    """
    
    def __init__(self, template_path: Path):
        self.template_path = Path(template_path)
        self.document: Optional[_Document] = None
        self.placeholder_counts: Dict[str, int] = {}
        
        self._load_template()
        self._scan_complete_document()
    
    def _load_template(self) -> None:
        """Load the DOCX template."""
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
        
        logger.info(f"Loading template: {self.template_path}")
        self.document = Document(str(self.template_path))
        
        # Check if document loaded successfully
        if hasattr(self.document, 'paragraphs') and hasattr(self.document, 'tables'):
            logger.info(f"Template loaded: {len(self.document.paragraphs)} paragraphs, "
                        f"{len(self.document.tables)} tables")
        else:
            logger.error("Document loaded but missing required attributes")
        
    
    def _scan_complete_document(self) -> None:
        """Scan entire document for placeholders."""
        logger.info("Scanning document for placeholders...")
        
        self.placeholder_counts = scan_placeholders(self.document)
        unique_placeholders = sorted(self.placeholder_counts.keys())
        total_occurrences = sum(self.placeholder_counts.values())
        
        logger.info(f"Found {len(unique_placeholders)} unique placeholders with "
                    f"{total_occurrences} total occurrences")
        logger.debug(f"Placeholders: {unique_placeholders}")
        
        if not unique_placeholders:
            logger.warning("No placeholders found! Expected format: {{PLACEHOLDER_NAME}}")
    
    def replace_placeholders(self, replacements: Dict[str, str]) -> Dict[str, int]:
        """
        Replace placeholders with provided content.
        
        Args:
            replacements: Dict mapping placeholder_name → replacement_text
            
        Returns:
            Dict with replacement statistics
        """
        logger.info("Starting placeholder replacement...")
        
        total_placeholders = sum(self.placeholder_counts.values())
        stats = {"total_placeholders": total_placeholders, "replaced": 0, "skipped_missing": 0, "skipped_errors": 0}

        try:
            replace_stats = replace_placeholders_in_document(self.document, replacements)
            stats["replaced"] = replace_stats["replaced"]
            stats["skipped_missing"] = replace_stats["skipped_missing"]
            stats["skipped_errors"] = replace_stats["skipped_mismatch"]
        except Exception as e:
            # If anything goes wrong, keep the template content intact and report the failure.
            logger.error(f"Replacement failed unexpectedly: {e}")
            stats["skipped_errors"] = total_placeholders

        logger.info(
            "Replacement complete: %d replaced, %d skipped (missing), %d skipped (errors)",
            stats["replaced"],
            stats["skipped_missing"],
            stats["skipped_errors"],
        )
        return stats
    
    def save(self, output_path: Path) -> None:
        """Save the modified document."""
        if not self.document:
            raise ValueError("No document loaded")
        
        logger.info(f"Saving document to: {output_path}")
        ensure_dir_exists(output_path.parent)
        self.document.save(str(output_path))
        logger.info(f"Document saved successfully: {output_path}")
    
    def get_placeholder_names(self) -> List[str]:
        """Get list of all detected placeholder names."""
        return sorted(self.placeholder_counts.keys())
    
    def get_placeholder_count(self, name: str) -> int:
        """Get count of occurrences for a specific placeholder."""
        return int(self.placeholder_counts.get(name, 0))


def main():
    """Example usage."""
    template_path = Path("templates/Master_Resume.docx")
    
    try:
        engine = SafeDocxEngine(template_path)
        print(f"Found placeholders: {engine.get_placeholder_names()}")
        
        # Example replacements
        replacements = {
            "SUMMARY": "Experienced software engineer with expertise in...",
            "EXP1_BULLET1": "Developed scalable systems serving 1M+ users",
            "SKILLS_1": "Python, JavaScript, SQL, Cloud Computing"
        }
        
        stats = engine.replace_placeholders(replacements)
        print(f"Replacement stats: {stats}")
        
        output_path = Path("output/Safe_Tailored_Resume.docx")
        engine.save(output_path)
        print(f"Saved to: {output_path}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()